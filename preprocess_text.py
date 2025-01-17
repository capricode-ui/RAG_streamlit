def preprocess_text(word_file_path, size):

  import os


  # Step 2: Read the Word file if it exists
  from docx import Document

  # Path to the Word file in the Kaggle dataset


  # Read the Word file
  doc = Document(word_file_path)

  # Extract text from the document
  full_text = []
  for para in doc.paragraphs:
      full_text.append(para.text)

  # Join all paragraphs into a single string
  document_text = '\n'.join(full_text)
  #print(document_text)
  from langchain.text_splitter import RecursiveCharacterTextSplitter
  from langchain.docstore.document import Document

  text_splitter = RecursiveCharacterTextSplitter(chunk_size=size, chunk_overlap=200)
  # Convert the string to a Document object
  docs = [Document(page_content=document_text)]
  text = text_splitter.split_documents(docs) # Pass the Document object to split_documents
  #/kaggle/input/dataset1-fd/FD.docx
  return text