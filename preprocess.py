import streamlit as st
def preprocess_text(files, links, size, overlap):
    import os
    from PyPDF2 import PdfReader
    from docx import Document as DocxDocument
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.docstore.document import Document
    from selenium import webdriver
    from selenium.webdriver.chrome.service import Service
    from selenium.webdriver.common.by import By
    import time

    paragraphs = []

    # Step 1: Process each file
    for file in files:
        if file.name.endswith(".pdf"):
            reader = PdfReader(file)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:  # Ensure the page has text
                    paragraphs.extend(page_text.split("\n"))  # Split by line breaks
        elif file.name.endswith(".docx"):
            docx = DocxDocument(file)
            for paragraph in docx.paragraphs:
                if paragraph.text.strip():  # Ensure the paragraph has text
                    paragraphs.append(paragraph.text)

    # Step 2: Process each link using Selenium
    chrome_driver_path = r"D:\chromedriver-win64\chromedriver-win64\chromedriver.exe"  # Update this path
    service = Service(chrome_driver_path)
    driver = webdriver.Chrome(service=service)

    for link in links:
        try:
            driver.get(link)
            time.sleep(3)  # Allow the page to load
            body_text = driver.find_element(By.TAG_NAME, "body").text
            paragraphs.extend(body_text.split("\n"))  # Add the page's text content

            # Extract FAQs
            try:
                faq_container = driver.find_element(By.CSS_SELECTOR, ".faqs.aem-GridColumn.aem-GridColumn--default--12")

                # Show more content if the button is available
                while True:
                    try:
                        show_more_button = faq_container.find_element(By.CSS_SELECTOR, ".accordion_toggle_show-more")
                        if show_more_button.is_displayed():
                            show_more_button.click()
                            time.sleep(1)
                        else:
                            break
                    except Exception:
                        break

                # Extract FAQ questions and answers
                toggle_buttons = faq_container.find_elements(By.CSS_SELECTOR, ".accordion_toggle, .accordion_row")
                all_faqs = []
                for button in toggle_buttons:
                    try:
                        button.click()
                        time.sleep(1)
                        expanded_content = faq_container.find_elements(By.CSS_SELECTOR, ".accordion_body, .accordionbody_links, .aem-rte-content")
                        for content in expanded_content:
                            text = content.text.strip()
                            if text and text not in [faq['answer'] for faq in all_faqs]:
                                question = button.text.strip()
                                if question:
                                    all_faqs.append({"question": question, "answer": text})

                    except Exception as e:
                        print(f"Error interacting with button: {e}")

                print("Entire Page Content:")
                print(body_text)

                print("\nExtracted FAQ Questions and Answers:")
                for i, faq in enumerate(all_faqs, start=1):
                    print(f"Q: {faq['question']}\n   A: {faq['answer']}\n")

            except Exception as faq_error:
                print(f"FAQ extraction failed for {link}: {faq_error}")
        except Exception as link_error:
            print(f"Failed to process link {link}: {link_error}")
        finally:
            driver.quit()

    # Step 3: Remove empty paragraphs
    paragraphs = [para.strip() for para in paragraphs if para.strip()]

    # Step 4: Convert paragraphs to Document objects
    docs = [Document(page_content=para) for para in paragraphs]

    # Step 5: Use RecursiveCharacterTextSplitter for chunking
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=size, chunk_overlap=overlap)
    text_chunks = text_splitter.split_documents(docs)

    return text_chunks



def preprocess_chroma(text, embedding_model_name, persist_directory):
    from langchain.embeddings import SentenceTransformerEmbeddings
    from langchain.vectorstores import Chroma

    embedding_model = SentenceTransformerEmbeddings(model_name=embedding_model_name)
    vectordb = Chroma.from_documents(documents=text, embedding=embedding_model, persist_directory=persist_directory)
    vectordb.persist()

    vectordb = Chroma(persist_directory=persist_directory, embedding_function=embedding_model)
    retriever = vectordb.as_retriever()

    return vectordb, retriever

def preprocess_faiss(text, embedding_model_name):
    from langchain.embeddings import SentenceTransformerEmbeddings
    import numpy as np
    import faiss
    from langchain.docstore.in_memory import InMemoryDocstore
    from langchain.vectorstores import FAISS
    from langchain.docstore.document import Document

    embedding_model = SentenceTransformerEmbeddings(model_name=embedding_model_name)
    texts = [doc.page_content for doc in text]
    embeddings = embedding_model.embed_documents(texts)
    embeddings = np.array(embeddings)
    dimension = embeddings.shape[1]

    index = faiss.IndexFlatL2(dimension)
    index.add(embeddings)

    docstore = InMemoryDocstore({i: Document(page_content=texts[i]) for i in range(len(texts))})
    index_to_docstore_id = {i: i for i in range(len(texts))}

    vector_store = FAISS(
        index=index,
        docstore=docstore,
        index_to_docstore_id=index_to_docstore_id,
        embedding_function=embedding_model.embed_query
    )

    return index, docstore, index_to_docstore_id, vector_store

def preprocess_qdrant(text, embeddings, client_url, client_api_key, collection_name, batch_size=250):
    from qdrant_client import QdrantClient, models

    client = QdrantClient(url=client_url, api_key=client_api_key)
    client.recreate_collection(
        collection_name=collection_name,
        vectors_config=models.VectorParams(size=embeddings.shape[1], distance=models.Distance.COSINE)
    )

    qdrant_index = list(range(1, len(text) + 1))
    for i in range(0, len(text), batch_size):
        low_idx = min(i + batch_size, len(text))
        batch_of_ids = qdrant_index[i: low_idx]
        batch_of_embs = embeddings[i: low_idx]
        batch_of_payloads = [{'page_content': doc.page_content, 'metadata': doc.metadata} for doc in text[i: low_idx]]

        client.upsert(
            collection_name=collection_name,
            points=models.Batch(
                ids=batch_of_ids,
                vectors=batch_of_embs.tolist(),
                payloads=batch_of_payloads
            )
        )

    return client

def preprocess_pinecone(text,embedding_model_name):
    import numpy as np
    from langchain.embeddings import SentenceTransformerEmbeddings
    embedding_model= SentenceTransformerEmbeddings(model_name=embedding_model_name)
    # Extract the 'page_content' (text) from each Document object
    texts = [doc.page_content for doc in text]
    embeddings = embedding_model.embed_documents(texts)  # Pass the list of texts
    embeddings = np.array(embeddings)
    embeddings = embeddings.tolist()


    import pinecone
    from pinecone.grpc import PineconeGRPC as Pinecone
    from pinecone import ServerlessSpec
    import uuid

    # ... (your existing code) ...

    index_name = "test5"

    pinecone = Pinecone(
        api_key="pcsk_42Yw14_EaKdaMLiAJfWub3s2sEJYPW3jyXXjdCYkH8Mh8rD8wWJ3pS6oCCC9PGqBNuDTuf",
        environment="us-east-1"
    )
    # Check if the index exists
    indexes = pinecone.list_indexes().names()

    if index_name in indexes:
        pinecone.delete_index(index_name)

    pinecone.create_index(
      name=index_name,
      dimension=len(embeddings[0]),
      metric="cosine",
      spec=ServerlessSpec(
        cloud="aws",
        region="us-east-1"
      ),
      deletion_protection="disabled"
    )

    pinecone_index = pinecone.Index(index_name)

    upsert_data = []
    for i in range(len(texts)):
      upsert_data.append((str(uuid.uuid4()), embeddings[i], {"text": texts[i]}))

    # Upsert data in batches (adjust batch_size as needed)
    batch_size = 100  # Example batch size
    for i in range(0, len(upsert_data), batch_size):
        batch = upsert_data[i : i + batch_size]
        pinecone_index.upsert(vectors=batch)
    return pinecone_index
    # ... (rest of your upsert code) ...

#!pip install weaviate-client

import numpy as np
import os
from langchain.embeddings import SentenceTransformerEmbeddings
import weaviate
from weaviate.auth import AuthApiKey
import weaviate.classes.config as wvcc
from weaviate.collections import Collection

def preprocess_weaviate(text, embedding_model_name):
    from langchain.embeddings import SentenceTransformerEmbeddings
    import numpy as np
    embedding_model = SentenceTransformerEmbeddings(model_name=embedding_model_name)
    import os
    os.environ["WEAVIATE_URL"] = "https://jrstxmtsfe1p4sw1imvta.c0.asia-southeast1.gcp.weaviate.cloud/"
    os.environ["WEAVIATE_API_KEY"] = "YoJvj3julM7BALAdHUjWrXPSkE95HsdxWtHn"

    weaviate_url = os.environ["WEAVIATE_URL"]
    weaviate_api_key = os.environ["WEAVIATE_API_KEY"]
    import weaviate
    from weaviate.auth import AuthApiKey

    client = weaviate.connect_to_weaviate_cloud(
        cluster_url=weaviate_url,
        auth_credentials=AuthApiKey(weaviate_api_key),
    )
    from langchain_weaviate.vectorstores import WeaviateVectorStore # Import WeaviateVectorStore

    vs = WeaviateVectorStore.from_documents(
        documents=text,
        embedding=embedding_model,
        client=client
    )

    return vs


# ipython-input-12-291a5c8eb7ff
from langchain.embeddings import SentenceTransformerEmbeddings

# Declare embedding_model_global as a global variable
embedding_model_global = None

def preprocess_vectordbs(files,links, embedding_model_name, size, overlap):



    global embedding_model_global  # Declare embedding_model_global as global within the function

    text = preprocess_text(files,links, size,overlap)
    st.success("Preprocessing Text Complete!")
    persist_directory = 'db'
    # Assign the model directly, not the model name
    embedding_model_global = SentenceTransformerEmbeddings(model_name=embedding_model_name)

    # Process Chroma
    vectordb, retriever = preprocess_chroma(text, embedding_model_name, persist_directory) #embedding_model_name changed to embedding_model_global
    st.success("Preprocessing Chroma Complete!")
    # Process FAISS
    index, docstore, index_to_docstore_id, vector_store = preprocess_faiss(text, embedding_model_name) #embedding_model_name changed to embedding_model_global
    st.success("Preprocessing Faiss Complete!")
    # Process Qdrant UNMASK LATER
    # pinecone
    vs = preprocess_weaviate(text, embedding_model_name)
    st.success("Preprocessing Weaviate Complete!")

    pinecone_index = preprocess_pinecone(text, embedding_model_name)
    st.success("Preprocessing Pinecone Complete!")
    # process weaviate

    # return vs for weaviate
    #embeddings = vector_store.index.reconstruct_n(0, len(text))
    #client_url = "https://186e02e2-6d10-4b48-baf1-273a91f6c628.us-east4-0.gcp.cloud.qdrant.io"
    #client_api_key = "Wc7kgaf6hXuYIHppaAT87CUyVy5pwigwGaI3oufb3r3Xbcwdo9c_jw"
    #collection_name = "text_vectors"
    #client = preprocess_qdrant(text, embeddings, client_url, client_api_key, collection_name)
    st.success("Preprocessing Qdrant Complete!")



    return index, docstore, index_to_docstore_id, vector_store, retriever, pinecone_index,embedding_model_global,vs


